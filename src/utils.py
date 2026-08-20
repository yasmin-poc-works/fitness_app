import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


INLINE_PATTERN = re.compile(
    r"(\*\*.+?\*\*|__.+?__|`.+?`|\[[^\]]+\]\([^)]+\)|\*[^*]+\*|_[^_]+_)"
)
ORDERED_LIST_PATTERN = re.compile(r"^(\d+)\.\s+")
UNORDERED_LIST_PATTERN = re.compile(r"^[-*+]\s+")
TABLE_SEPARATOR_PATTERN = re.compile(r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?\s*$")


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_run(paragraph, text: str, bold: bool = False, italic: bool = False, code: bool = False) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if code:
        run.font.name = "Courier New"
        run.font.size = Pt(10)


def _parse_inline_markup(paragraph, text: str) -> None:
    """Render a subset of markdown inline formatting into a Word paragraph."""
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            _add_run(paragraph, text[cursor:match.start()])

        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            _add_run(paragraph, token[2:-2], bold=True)
        elif token.startswith("__") and token.endswith("__"):
            _add_run(paragraph, token[2:-2], bold=True)
        elif token.startswith("`") and token.endswith("`"):
            _add_run(paragraph, token[1:-1], code=True)
        elif token.startswith("[") and "](" in token and token.endswith(")"):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            _add_hyperlink(paragraph, label, url)
        elif token.startswith("*") and token.endswith("*"):
            _add_run(paragraph, token[1:-1], italic=True)
        elif token.startswith("_") and token.endswith("_"):
            _add_run(paragraph, token[1:-1], italic=True)
        else:
            _add_run(paragraph, token)
        cursor = match.end()

    if cursor < len(text):
        _add_run(paragraph, text[cursor:])


def _add_paragraph(doc: Document, text: str, style: str = "Normal", level: int | None = None):
    if level is not None:
        paragraph = doc.add_heading(text, level=level)
        return paragraph

    paragraph = doc.add_paragraph(style=style)
    _parse_inline_markup(paragraph, text)
    return paragraph


def _is_table_block(lines: list[str], index: int) -> bool:
    return (
        index + 1 < len(lines)
        and "|" in lines[index]
        and TABLE_SEPARATOR_PATTERN.match(lines[index + 1].strip()) is not None
    )


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def _add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.style = "Table Grid"

    for row_index, row_values in enumerate(rows):
        for col_index in range(len(table.columns)):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            if col_index >= len(row_values):
                continue
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _parse_inline_markup(paragraph, row_values[col_index])
            if row_index == 0:
                _set_cell_shading(cell, "D9EAF7")
                for run in paragraph.runs:
                    run.bold = True


def _parse_table_block(lines: list[str], start_index: int) -> tuple[list[list[str]], int]:
    rows = [_split_table_row(lines[start_index])]
    index = start_index + 2
    while index < len(lines) and "|" in lines[index].strip():
        rows.append(_split_table_row(lines[index]))
        index += 1
    return rows, index


def convert_to_docx(text: str) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = text.splitlines()
    index = 0
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()

        if not line:
            index += 1
            continue

        if line.startswith("```"):
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index].rstrip("\n"))
                index += 1
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            doc.add_heading(heading_match.group(2).strip(), level=level)
            index += 1
            continue

        if _is_table_block(lines, index):
            table_rows, next_index = _parse_table_block(lines, index)
            _add_markdown_table(doc, table_rows)
            index = next_index
            continue

        unordered_match = UNORDERED_LIST_PATTERN.match(line)
        ordered_match = ORDERED_LIST_PATTERN.match(line)
        if unordered_match:
            paragraph = doc.add_paragraph(style="List Bullet")
            _parse_inline_markup(paragraph, line[unordered_match.end():].strip())
            index += 1
            continue
        if ordered_match:
            paragraph = doc.add_paragraph(style="List Number")
            _parse_inline_markup(paragraph, line[ordered_match.end():].strip())
            index += 1
            continue

        paragraph = doc.add_paragraph(style="Normal")
        _parse_inline_markup(paragraph, line)
        index += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
